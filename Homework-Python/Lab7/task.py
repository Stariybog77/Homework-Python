#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from abc import ABC, abstractmethod
import PySimpleGUI as sg
from docx import Document


# =========================
# АБСТРАКТНЫЙ КЛАСС
# =========================
class Building(ABC):

    @abstractmethod
    def area(self):
        pass

    @abstractmethod
    def heat(self):
        pass


# =========================
# КОМНАТА
# =========================
class Room(Building):

    def __init__(self, length, width):
        self._length = length
        self._width = width

    @property
    def length(self):
        return self._length

    @property
    def width(self):
        return self._width

    def area(self):
        return self._length * self._width

    def heat(self):
        return self.area() * 100

    def __str__(self):
        return f"Комната {self._length}x{self._width}"

    def __len__(self):
        return int(self.area())


# =========================
# КВАРТИРА
# =========================
class Apartment(Building):

    def __init__(self, rooms):
        self._rooms = rooms

    @property
    def rooms(self):
        return self._rooms

    def area(self):
        return sum(room.area() for room in self._rooms)

    def heat(self):
        return sum(room.heat() for room in self._rooms)

    def __str__(self):
        return f"Квартира ({len(self._rooms)} комнаты)"

    def __len__(self):
        return len(self._rooms)


# =========================
# ДОМ
# =========================
class House(Building):

    def __init__(self, apartments):
        self._apartments = apartments

    @property
    def apartments(self):
        return self._apartments

    def area(self):
        return sum(ap.area() for ap in self._apartments)

    def heat(self):
        return sum(ap.heat() for ap in self._apartments)

    def __str__(self):
        return f"Дом ({len(self._apartments)} квартиры)"

    def __len__(self):
        return len(self._apartments)


# =========================
# GUI
# =========================
layout = [
    [sg.Text("Тип помещения")],
    [sg.Combo(["Комната", "Квартира", "Дом"], key="type")],

    [sg.Text("Длина"), sg.Input(key="length")],
    [sg.Text("Ширина"), sg.Input(key="width")],

    [sg.Button("Рассчитать")],
    [sg.Button("Сохранить в DOC")],

    [sg.Text("", key="result")]
]

window = sg.Window("Lab7", layout)

last_text = ""

while True:
    event, values = window.read()

    if event == sg.WIN_CLOSED:
        break

    if event == "Рассчитать":
        try:
            l = float(values["length"])
            w = float(values["width"])

            if values["type"] == "Комната":
                obj = Room(l, w)

            elif values["type"] == "Квартира":
                obj = Apartment([
                    Room(l, w),
                    Room(l + 1, w + 1)
                ])

            elif values["type"] == "Дом":
                obj = House([
                    Apartment([
                        Room(l, w),
                        Room(l + 1, w)
                    ]),
                    Apartment([
                        Room(l, w + 1),
                        Room(l + 2, w)
                    ])
                ])

            else:
                window["result"].update("Выберите тип помещения!")
                continue

            area = obj.area()
            heat = obj.heat()

            result_text = f"{obj}\nПлощадь: {area:.2f} м²\nТепло: {heat:.2f} Вт"
            window["result"].update(result_text)

            last_text = result_text

        except:
            window["result"].update("Ошибка ввода!")

    if event == "Сохранить в DOC":
        if not last_text:
            window["result"].update("Сначала сделайте расчёт!")
        else:
            doc = Document()
            doc.add_heading("Отчёт по расчёту", 0)
            doc.add_paragraph(last_text)

            doc.save("report.docx")
            window["result"].update("Файл report.docx сохранён!")

window.close()