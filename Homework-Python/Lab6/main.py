#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import tkinter as tk
from tkinter import messagebox
from docx import Document

from building.room import room_area, room_heat
from building.apartment import apartment_area, apartment_heat
from building.building import building_area, building_heat


# =========================
# РАСЧЁТ
# =========================
def calculate():
    try:
        length = float(entry_length.get())
        width = float(entry_width.get())

        if room_type.get() == "room":
            area = room_area(length, width)
            heat = room_heat(area)

        elif room_type.get() == "apartment":
            rooms = [(length, width), (length, width)]  # 2 комнаты
            area = apartment_area(rooms)
            heat = apartment_heat(rooms)

        elif room_type.get() == "building":
            apartments = [
                [(length, width), (length, width)],
                [(length, width), (length, width)]
            ]
            area = building_area(apartments)
            heat = building_heat(apartments)

        result_label.config(text=f"Площадь: {area:.2f} м²\nТепло: {heat:.2f} Вт")

        global last_result
        last_result = (area, heat)

    except:
        messagebox.showerror("Ошибка", "Введите корректные числа!")


# =========================
# СОХРАНЕНИЕ В DOCX
# =========================
def save_to_doc():
    if last_result is None:
        messagebox.showerror("Ошибка", "Сначала выполните расчёт!")
        return

    doc = Document()
    doc.add_heading('Отчёт по расчёту помещения', 0)

    doc.add_paragraph(f"Тип помещения: {room_type.get()}")
    doc.add_paragraph(f"Площадь: {last_result[0]:.2f} м²")
    doc.add_paragraph(f"Тепловая мощность: {last_result[1]:.2f} Вт")

    doc.save("report.docx")
    messagebox.showinfo("Готово", "Файл report.docx сохранён!")


# =========================
# GUI
# =========================
window = tk.Tk()
window.title("Расчёт помещения")
window.geometry("320x320")

last_result = None

# Выбор типа помещения
room_type = tk.StringVar(value="room")

tk.Label(window, text="Тип помещения:").pack()

tk.Radiobutton(window, text="Комната", variable=room_type, value="room").pack()
tk.Radiobutton(window, text="Квартира", variable=room_type, value="apartment").pack()
tk.Radiobutton(window, text="Дом", variable=room_type, value="building").pack()

# Ввод данных
tk.Label(window, text="Длина").pack()
entry_length = tk.Entry(window)
entry_length.pack()

tk.Label(window, text="Ширина").pack()
entry_width = tk.Entry(window)
entry_width.pack()

# Кнопки
tk.Button(window, text="Рассчитать", command=calculate).pack(pady=5)
tk.Button(window, text="Сохранить в DOC", command=save_to_doc).pack(pady=5)

# Результат
result_label = tk.Label(window, text="")
result_label.pack(pady=10)

window.mainloop()